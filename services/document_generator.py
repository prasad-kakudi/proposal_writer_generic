from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from datetime import datetime

class DocumentGenerator:
    """Service for generating DOCX documents"""
    
    def __init__(self):
        pass
    
    def create_docx(self, content: str, output_path: str) -> str:
        """Generate DOCX document from content"""
        try:
            # Create new document
            doc = Document()
            
            # Add header
            self._add_header(doc)
            
            # Add content
            self._add_content(doc, content)
            
            # Add footer
            self._add_footer(doc)
            
            # Save document
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error creating DOCX document: {str(e)}")
    
    def _add_header(self, doc: Document):
        """Add professional header to document"""
        # Title
        title = doc.add_heading('Request for Proposal Response', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Date
        date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add separator
        doc.add_paragraph('=' * 60).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()  # Empty line
    
    def _add_content(self, doc: Document, content: str):
        """Add main content to document"""
        # Split content into sections if possible
        sections = self._parse_content_sections(content)
        
        for section in sections:
            if section['type'] == 'heading':
                heading = doc.add_heading(section['text'], level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif section['type'] == 'paragraph':
                doc.add_paragraph(section['text'])
            elif section['type'] == 'bullet':
                p = doc.add_paragraph(section['text'], style='List Bullet')
    
    def _parse_content_sections(self, content: str) -> list:
        """Parse content into structured sections"""
        sections = []
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a heading (contains certain keywords or is short and capitalized)
            if (len(line) < 100 and 
                (line.isupper() or 
                 any(keyword in line.lower() for keyword in ['overview', 'executive summary', 'introduction', 'conclusion', 'requirements', 'approach', 'methodology', 'timeline', 'budget', 'team']))):
                sections.append({'type': 'heading', 'text': line})
            # Check if line is a bullet point
            elif line.startswith(('•', '-', '*', '1.', '2.', '3.', '4.', '5.')):
                sections.append({'type': 'bullet', 'text': line})
            else:
                sections.append({'type': 'paragraph', 'text': line})
        
        return sections
    
    def _add_footer(self, doc: Document):
        """Add footer to document"""
        doc.add_paragraph()  # Empty line
        doc.add_paragraph('=' * 60).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        footer_text = doc.add_paragraph('This document was generated using AI-powered RFP Response System')
        footer_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Make footer text smaller and italic
        run = footer_text.runs[0]
        run.font.size = Inches(0.1)
        run.italic = True
